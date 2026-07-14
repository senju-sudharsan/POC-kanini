from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT = Path(__file__).with_name("Data_Engineering_PoC_Documentation.docx")
FIGURES = Path(__file__).with_name("figures")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(paragraph, before=0, after=6, line=1.1, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_text(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph(p)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, after=3, line=1.1)
    set_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_paragraph(p, before=16 if level == 1 else 10, after=6, line=1.0)
    set_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, header_text in enumerate(headers):
        cell = header.cells[index]
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        set_font(p.add_run(header_text), size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_width(cells[index], widths[index])
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[index].paragraphs[0]
            set_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(str(value)), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_flow(doc, labels):
    table = doc.add_table(rows=1, cols=len(labels) * 2 - 1)
    table.autofit = False
    for index, label in enumerate(labels):
        cell = table.cell(0, index * 2)
        set_cell_width(cell, 1.02)
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_font(p.add_run(label), size=8.5, color=DARK_BLUE, bold=True)
        if index < len(labels) - 1:
            arrow = table.cell(0, index * 2 + 1)
            set_cell_width(arrow, 0.08)
            p = arrow.paragraphs[0]
            set_paragraph(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_font(p.add_run("→"), size=11, color=BLUE, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _diagram_font(size, bold=False):
    font_name = "C:\\Windows\\Fonts\\segoeuib.ttf" if bold else "C:\\Windows\\Fonts\\segoeui.ttf"
    return ImageFont.truetype(font_name, size)


def _color(value):
    return f"#{value}" if len(value) == 6 and not value.startswith("#") else value


def _draw_centered(draw, box, text, font, fill="FFFFFF"):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=4, align="center")
    x = left + (right - left - (bbox[2] - bbox[0])) / 2
    y = top + (bottom - top - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), text, font=font, fill=_color(fill), spacing=4, align="center")


def _box(draw, box, text, fill, outline, size=25):
    draw.rounded_rectangle(box, radius=18, fill=_color(fill), outline=_color(outline), width=4)
    _draw_centered(draw, box, text, _diagram_font(size, bold=True))


def _arrow(draw, start, end, color="A855F7", width=6):
    draw.line([start, end], fill=_color(color), width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 18 * direction, y2 - 11), (x2 - 18 * direction, y2 + 11)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 11, y2 - 18 * direction), (x2 + 11, y2 - 18 * direction)]
    draw.polygon(points, fill=_color(color))


def make_diagrams():
    FIGURES.mkdir(exist_ok=True)
    background = "#0F172A"

    # Diagram 1: supplied end-to-end architecture.
    image = Image.new("RGB", (1600, 980), background)
    draw = ImageDraw.Draw(image)
    _draw_centered(draw, (0, 28, 1600, 80), "End-to-End Data Engineering Architecture", _diagram_font(34, True), "FFFFFF")
    stages = [
        ("CSV Files", "1E293B", "A855F7"), ("Python\nIngestion", "7C3AED", "A855F7"),
        ("Bronze\nLayer", "3B2F12", "D4A017"), ("Silver\nLayer", "334155", "CBD5E1"),
        ("Gold\nLayer", "4A3B12", "FFD700"), ("Validation\nFramework", "064E3B", "10B981"),
        ("FastAPI", "7C3AED", "A855F7"), ("React\nDashboard", "A855F7", "C084FC"),
    ]
    x, y, w, h = 70, 230, 155, 135
    for idx, (label, fill, stroke) in enumerate(stages):
        left = x + idx * 195
        _box(draw, (left, y, left + w, y + h), label, fill, stroke, 22)
        if idx < len(stages) - 1:
            _arrow(draw, (left + w, y + h / 2), (left + 195, y + h / 2))
    _box(draw, (290, 560, 580, 680), "Metadata Framework\nbatch_control", "5B21B6", "C084FC", 22)
    _box(draw, (690, 560, 910, 680), "Airflow DAG", "0F172A", "38BDF8", 22)
    _arrow(draw, (430, 560), (430, 380), "C084FC", 5)
    for target_x in (360, 555, 750, 945):
        _arrow(draw, (800, 560), (target_x, 380), "38BDF8", 4)
    _draw_centered(draw, (0, 805, 1600, 890), "CSV -> Bronze -> Silver -> Gold -> Validation -> API -> Dashboard", _diagram_font(24), "CBD5E1")
    image.save(FIGURES / "end_to_end_architecture.png")

    # Diagram 2: supplied warehouse lineage.
    image = Image.new("RGB", (1600, 1040), background)
    draw = ImageDraw.Draw(image)
    _draw_centered(draw, (0, 22, 1600, 72), "Olist Warehouse Lineage and Consumption", _diagram_font(34, True), "FFFFFF")
    _box(draw, (35, 390, 205, 540), "Olist CSV\nFiles", "1E293B", "A855F7", 24)
    groups = [
        (260, "Bronze Layer", "3B2F12", "D4A017", ["customers_raw", "orders_raw", "products_raw", "payments_raw", "sellers_raw"]),
        (610, "Silver Layer", "334155", "CBD5E1", ["customers", "orders", "products", "payments", "sellers", "order_fact", "customers_scd"]),
        (980, "Gold Layer", "4A3B12", "FFD700", ["sales_summary", "product_performance", "seller_performance"]),
    ]
    for left, title, fill, stroke, tables in groups:
        draw.rounded_rectangle((left, 160, left + 300, 820), radius=22, fill=_color(fill), outline=_color(stroke), width=4)
        _draw_centered(draw, (left, 178, left + 300, 228), title, _diagram_font(25, True))
        for index, table in enumerate(tables):
            top = 250 + index * 75
            draw.rounded_rectangle((left + 30, top, left + 270, top + 52), radius=10, fill=_color("1E293B"), outline=_color(stroke), width=2)
            _draw_centered(draw, (left + 30, top, left + 270, top + 52), table, _diagram_font(17), "FFFFFF")
    _box(draw, (1350, 320, 1510, 430), "FastAPI", "7C3AED", "A855F7", 21)
    _box(draw, (1350, 570, 1510, 680), "React\nDashboard", "A855F7", "C084FC", 21)
    _arrow(draw, (205, 465), (260, 465))
    _arrow(draw, (560, 465), (610, 465), "CBD5E1")
    _arrow(draw, (910, 465), (980, 465), "FFD700")
    _arrow(draw, (1280, 465), (1350, 375), "A855F7")
    _arrow(draw, (1430, 430), (1430, 570), "C084FC")
    _draw_centered(draw, (0, 890, 1600, 950), "Raw source preservation, standardized conformance, analytics aggregation, and governed dashboard delivery", _diagram_font(21), "CBD5E1")
    image.save(FIGURES / "warehouse_lineage.png")

    # Diagram 3: supplied SCD Type 2 flow.
    image = Image.new("RGB", (1600, 620), background)
    draw = ImageDraw.Draw(image)
    _draw_centered(draw, (0, 25, 1600, 75), "SCD Type 2 Customer Location Change", _diagram_font(34, True), "FFFFFF")
    flow = [
        ("Version 1\nOSASCO\nCurrent", "334155", "94A3B8"),
        ("Change\nDetected", "7C3AED", "A855F7"),
        ("Expire Version 1\nis_current = false", "991B1B", "EF4444"),
        ("Insert\nVersion 2", "7C3AED", "A855F7"),
        ("Version 2\nDEMO_CITY_A\nCurrent", "A855F7", "C084FC"),
    ]
    x, y, w, h = 50, 220, 245, 150
    for index, (label, fill, stroke) in enumerate(flow):
        left = x + index * 315
        _box(draw, (left, y, left + w, y + h), label, fill, stroke, 21)
        if index < len(flow) - 1:
            _arrow(draw, (left + w, y + h / 2), (left + 315, y + h / 2))
    _draw_centered(draw, (0, 470, 1600, 535), "Historical row retained with an end date; a new current row receives the next version number.", _diagram_font(22), "CBD5E1")
    image.save(FIGURES / "scd_type_2_flow.png")


def add_figure(doc, image_path, caption):
    doc.add_picture(str(image_path), width=Inches(6.55))
    p = doc.add_paragraph()
    set_paragraph(p, after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run(caption), size=9, color="6B7280", italic=True)


def add_code(doc, title, code):
    p = doc.add_paragraph()
    set_paragraph(p, before=6, after=3)
    set_font(p.add_run(title), size=10, color=DARK_BLUE, bold=True)
    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.2)
    code_p.paragraph_format.space_after = Pt(6)
    code_p.paragraph_format.line_spacing = 1.0
    run = code_p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(INK)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    code_p._p.get_or_add_pPr().append(shading)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color="6B7280")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build():
    make_diagrams()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Cover page
    doc.add_paragraph().paragraph_format.space_after = Pt(80)
    p = doc.add_paragraph()
    set_paragraph(p, after=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("PROJECT DOCUMENTATION"), size=11, color=BLUE, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("Data Engineering PoC - Olist E-Commerce Data Warehouse"), size=27, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, after=40, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("End-to-End Data Engineering Pipeline using Medallion Architecture"), size=14, color=INK)
    add_table(doc, ["Prepared by", "Internship / Organization", "Completion Date"], [["Sudharsan S A", "KANINI", "14 / 07 / 2026"]], [2.15, 2.15, 2.2])
    p = doc.add_paragraph()
    set_paragraph(p, before=72, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("PostgreSQL | Docker | Apache Airflow | FastAPI | React"), size=11, color="6B7280")
    doc.add_page_break()

    add_heading(doc, "1. Executive Summary")
    add_text(doc, "This project implements an end-to-end data engineering proof of concept for the Olist e-commerce dataset. It ingests CSV source data into PostgreSQL, applies a Bronze-Silver-Gold Medallion Architecture, validates warehouse outputs, exposes analytics through FastAPI, and presents the results in a React dashboard.")
    add_text(doc, "The solution is designed to demonstrate practical warehouse engineering: raw data traceability, standardized Silver datasets, Gold analytics tables, batch metadata, incremental loading, validation checks, Apache Airflow orchestration, and customer Slowly Changing Dimension (SCD) processing.")
    add_table(doc, ["Objective", "Delivered capability"], [
        ["Reliable data pipeline", "CSV ingestion, batch status tracking, and Docker-hosted PostgreSQL."],
        ["Analytics-ready warehouse", "Gold revenue, product, and seller performance tables."],
        ["Operational governance", "Metadata control, validation reporting, incremental customer loading, and SCD history."],
        ["Business consumption", "FastAPI endpoints and an executive React dashboard."],
    ], [2.0, 4.5])

    add_heading(doc, "2. Project Architecture")
    add_text(doc, "The architecture separates ingestion, conformance, analytics, validation, and presentation responsibilities. PostgreSQL stores the warehouse schemas; Docker provides the database runtime; Airflow coordinates the warehouse stages; FastAPI serves governed API responses; and React consumes those responses for dashboards.")
    add_flow(doc, ["CSV Files", "Bronze", "Silver", "Gold", "Validation", "FastAPI", "React Dashboard"])
    add_figure(doc, FIGURES / "end_to_end_architecture.png", "Figure 1. Supplied end-to-end architecture diagram: ingestion, metadata, orchestration, validation, API, and dashboard flow.")
    add_table(doc, ["Component", "Role in the solution"], [
        ["Docker + PostgreSQL", "Runs the de_poc database in a reproducible containerized environment."],
        ["Bronze / Silver / Gold", "Implements raw capture, conformance, and analytics aggregation layers."],
        ["Apache Airflow", "Runs ingestion, Silver transformations, Gold transformations, and validation as a DAG."],
        ["FastAPI", "Provides overview, pipeline, quality, medallion, analytics, and SCD endpoints."],
        ["React", "Renders Overview, Dataset Explorer, Pipeline Health, Data Quality, and Business Intelligence views."],
    ], [2.0, 4.5])

    add_heading(doc, "3. Technology Stack")
    add_table(doc, ["Technology", "Purpose"], [
        ["Python", "Pipeline scripts, ingestion utilities, validation, SCD processing, and FastAPI services."],
        ["PostgreSQL", "Warehouse database (de_poc) with bronze, silver, gold, and metadata schemas."],
        ["Docker", "Reproducible PostgreSQL 16 runtime and persistent postgres_data volume."],
        ["SQLAlchemy", "Database engine and transactional batch control during CSV ingestion."],
        ["Pandas", "Reads source CSV files and loads raw Bronze tables."],
        ["FastAPI", "REST API layer with a consistent data/meta response envelope."],
        ["React + TypeScript", "Dashboard frontend built with React Query, Recharts, and Tailwind CSS."],
        ["Apache Airflow", "Orchestrates the warehouse_pipeline DAG."],
    ], [1.45, 5.05])

    add_heading(doc, "4. Dataset Overview")
    add_text(doc, "The Olist e-commerce dataset represents marketplace activity across customers, orders, products, sellers, payments, reviews, and geolocation. The project uses the CSV extracts as the source system and preserves batch/source metadata during ingestion.")
    add_table(doc, ["Source file / entity", "Warehouse use"], [
        ["customers", "Customer identity and city/state attributes; source for customer SCD processing."],
        ["orders", "Order lifecycle timestamps, status, and customer relationship."],
        ["order_items", "Item-level price, freight, product, and seller records."],
        ["payments", "Payment type, installment, and payment value."],
        ["products", "Product catalog and category attributes."],
        ["sellers", "Seller identity and location data."],
        ["reviews", "Customer review score and comments for future experience analysis."],
        ["geolocation", "State, city, latitude, longitude, and postal-prefix reference data."],
    ], [2.0, 4.5])

    add_heading(doc, "5. Medallion Architecture")
    add_heading(doc, "Bronze Layer", 2)
    add_text(doc, "Bronze preserves raw ingested records, including batch_id, source_system, and load timestamps. Representative tables include bronze.customers_raw, bronze.orders_raw, bronze.order_items_raw, bronze.payments_raw, bronze.products_raw, bronze.sellers_raw, bronze.reviews_raw, bronze.geolocation_raw, and bronze.category_translation_raw.")
    add_heading(doc, "Silver Layer", 2)
    add_text(doc, "Silver standardizes and cleans operational data. Core tables are silver.customers, silver.orders, silver.products, silver.sellers, silver.payments, and silver.order_fact. Customer city and state are normalized to uppercase, and order_fact joins order headers with item records. silver.customers_scd adds Type 1 and Type 2 customer history support.")
    add_heading(doc, "Gold Layer", 2)
    add_text(doc, "Gold exposes business-ready aggregations: gold.sales_summary, gold.product_performance, and gold.seller_performance. These tables support revenue trends, category rankings, and seller leaderboards in the dashboard.")
    add_flow(doc, ["Raw CSV", "bronze.*_raw", "silver.*", "gold.*"])
    add_figure(doc, FIGURES / "warehouse_lineage.png", "Figure 2. Supplied warehouse lineage diagram: Olist source files through Bronze, Silver, Gold, FastAPI, and the React dashboard.")

    add_heading(doc, "6. ETL Pipeline")
    add_text(doc, "The pipeline follows CSV -> Bronze -> Silver -> Gold. Ingestion scripts create a metadata.batch_control record, read CSV data through Pandas, enrich it with source and batch metadata, append it to the Bronze schema, and update the batch status when complete.")
    add_table(doc, ["Capability", "Implementation in repository"], [
        ["Batch processing", "metadata.batch_control stores pipeline_name, status, timing, processed records, and rejected records."],
        ["Incremental loading", "scripts/incremental_load.py reads metadata.load_tracker, selects Bronze customer rows newer than last_batch_id, then updates the tracker."],
        ["Silver transformation", "scripts/silver_transformations.py executes 01_customers.sql through 06_order_fact.sql."],
        ["Gold transformation", "scripts/gold_transformations.py executes sales, product, and seller aggregation SQL."],
        ["Audit tracking", "Pipeline and validation endpoints expose batch metadata, records processed, duration, and validation status."],
    ], [2.0, 4.5])

    add_heading(doc, "7. Data Quality and Validation")
    add_text(doc, "The validation script checks row availability across the conformed and analytics layers: Silver customers, orders, products, sellers, payments, order_fact, and the three Gold tables. API quality endpoints additionally expose row counts, integrity checks, validation results, and a quality score.")
    add_table(doc, ["Validation example", "Purpose"], [
        ["Record counts", "Confirms that expected Silver and Gold datasets contain rows after transformation."],
        ["Null / completeness checks", "Supported by the quality validation framework for operational field quality."],
        ["Integrity checks", "Checks duplicate customers and customer/order referential integrity."],
        ["Reconciliation checks", "Compares Bronze and Silver entity row counts for customer, order, product, seller, and payment entities."],
    ], [2.0, 4.5])
    add_text(doc, "Data Quality Score = 100 when the latest batch completes successfully with zero rejected records, as exposed through the quality and analytics dashboard posture.", bold_prefix="Data Quality Score = 100")

    add_heading(doc, "8. Slowly Changing Dimensions")
    add_text(doc, "The project adds silver.customers_scd for customer city/state dimension management. It stores customer_id, city, state, effective start/end dates, is_current, version_number, and created_at. A partial unique index ensures only one current version per customer.")
    add_table(doc, ["SCD approach", "Behavior"], [
        ["SCD Type 1", "Overwrites the current customer_city and customer_state values; no additional history version is created."],
        ["SCD Type 2", "Detects city/state changes, expires the current row, records effective_end_date, sets is_current to false, and inserts the next version."],
    ], [1.55, 4.95])
    add_text(doc, "Demonstrated customer: 00012a2ce6f8dcda20d059ce98491703")
    add_flow(doc, ["Version 1", "Change Detected", "Version 2"])
    add_table(doc, ["Version", "City", "Current state", "History treatment"], [
        ["1", "OSASCO", "Historical", "Preserved with effective_end_date and is_current = false."],
        ["2", "DEMO_CITY_A", "Current", "Inserted as the active row with a new effective_start_date."],
    ], [0.65, 1.3, 1.15, 3.4])
    add_figure(doc, FIGURES / "scd_type_2_flow.png", "Figure 3. Supplied SCD Type 2 diagram showing the OSASCO to DEMO_CITY_A customer location change.")
    add_text(doc, "The FastAPI SCD endpoints are GET /api/v1/scd/summary and GET /api/v1/scd/history/{customer_id}. The dashboard uses the history response to render the version timeline.")

    add_heading(doc, "9. Airflow Orchestration")
    add_text(doc, "airflow/dags/warehouse_pipeline.py defines the warehouse_pipeline DAG with four BashOperator tasks: bronze_ingestion, silver_transformations, gold_transformations, and validation. The DAG is manual (schedule=None), starts on 2026-07-01, and executes the stages sequentially.")
    add_flow(doc, ["bronze_ingestion", "silver_transformations", "gold_transformations", "validation"])
    add_text(doc, "Airflow provides a single operational view for pipeline execution, task status, and failure monitoring while the batch_control table provides the persisted warehouse metadata.")

    add_heading(doc, "10. Dashboard Overview")
    add_table(doc, ["Dashboard area", "Purpose"], [
        ["Overview", "Warehouse totals and a visible project feature inventory."],
        ["Dataset Explorer", "Bronze, Silver, Gold, and Business Transformations cards with table-level lineage and detail."],
        ["Pipeline Health", "Latest batch status, duration, rows processed, validation, capability cards, history, and batch summary."],
        ["Data Quality", "Quality score, row counts, integrity checks, and validation result visibility."],
        ["Business Intelligence", "Revenue trend, product categories, seller leaderboard, payment mix, funnel, geographic treemap, and SCD timeline."],
    ], [1.65, 4.85])
    add_text(doc, "The dashboard is backed by real API data. Key analytics endpoints include /api/v1/analytics/revenue-trend, /top-categories, /seller-performance, and /payment-distribution; SCD views use the two SCD endpoints listed above.")

    add_heading(doc, "11. Sample Code Snippets")
    add_code(doc, "Bronze ingestion - create a batch record (scripts/ingestion/load_customers.py)", """result = conn.execute(text(\"\"\"
    INSERT INTO metadata.batch_control (pipeline_name, status)
    VALUES ('customers_ingestion', 'RUNNING')
    RETURNING batch_id
\"\"\"))
batch_id = result.scalar()""")
    add_code(doc, "Silver transformation - standardize customer location (sql/silver/01_customers.sql)", """SELECT DISTINCT ON (customer_id)
    customer_id, customer_unique_id,
    UPPER(customer_city), UPPER(customer_state), batch_id
FROM bronze.customers_raw
ORDER BY customer_id, batch_id DESC;""")
    add_code(doc, "Gold aggregation - daily revenue (sql/gold/01_sales_summary.sql)", """SELECT DATE(purchase_timestamp),
       COUNT(DISTINCT order_id), SUM(price), SUM(freight_value)
FROM silver.order_fact
GROUP BY DATE(purchase_timestamp);""")
    add_code(doc, "SCD Type 2 - expire current row (scripts/scd.py)", """UPDATE silver.customers_scd AS target
SET effective_end_date = CURRENT_TIMESTAMP, is_current = FALSE
FROM silver.customers AS source
WHERE target.customer_id = source.customer_id
  AND target.is_current = TRUE
  AND (target.customer_city IS DISTINCT FROM source.customer_city
       OR target.customer_state IS DISTINCT FROM source.customer_state);""")
    add_code(doc, "Validation - verify a transformed table (scripts/validation.py)", """VALIDATIONS = [
    ('Silver Customers', 'SELECT COUNT(*) FROM silver.customers'),
    ('Gold Sales Summary', 'SELECT COUNT(*) FROM gold.sales_summary'),
]""")
    add_code(doc, "FastAPI endpoint - customer SCD history (backend/app/routes/scd.py)", """@router.get('/history/{customer_id}')
def get_history(customer_id: str):
    history = get_customer_scd_history(customer_id)
    if history is None:
        raise HTTPException(status_code=404)
    return success_response(history.model_dump())""")

    add_heading(doc, "12. Challenges and Solutions")
    add_table(doc, ["Challenge", "Applied solution"], [
        ["PostgreSQL connectivity", "Centralized connection helpers and Docker environment variables provide consistent access to de_poc."],
        ["Docker networking", "Pipeline scripts use the de_poc_postgres container hostname within the Docker/Airflow environment."],
        ["Incremental updates", "metadata.load_tracker records the last processed customer batch and limits the next load to newer Bronze rows."],
        ["Customer attribute changes", "SCD Type 2 expires the existing current record and inserts a versioned replacement."],
        ["Dashboard integration", "React Query services consume FastAPI's consistent data/meta envelope and display loading, empty, and error states."],
    ], [2.0, 4.5])

    add_heading(doc, "13. Outcomes")
    add_text(doc, "The proof of concept demonstrates a complete data engineering path from e-commerce CSV data to governed warehouse analytics and a presentation-ready dashboard.")
    for item in [
        "Bronze Layer", "Silver Layer", "Gold Layer", "Incremental Loading", "Metadata Framework",
        "Audit Framework", "Validation Framework", "SCD Type 1", "SCD Type 2", "Airflow", "FastAPI APIs", "Dashboard",
    ]:
        add_bullet(doc, f"Implemented: {item}")
    add_text(doc, "The resulting platform provides traceable ingestion, clean and reusable dimensional data, business metrics, observable pipeline status, and a concrete demonstration of customer history management.")

    # Ensure all pages use the same footer after page breaks.
    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        if not footer.text:
            add_page_number(footer)

    doc.core_properties.author = "Sudharsan S A"
    doc.core_properties.title = "Data Engineering PoC - Olist E-Commerce Data Warehouse"
    doc.core_properties.subject = "Internship project documentation"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
